from docx import Document
import os

from config import output_dir

def iter_headings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph

def iter_text(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Normal'):
            yield paragraph

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

def split_word(filename: str):
    i = 0
    current_part_name = ""
    current_document = Document()

    document = Document(filename)
    basename = os.path.basename(filename)

    for p in document.paragraphs:
        if p.style.name.startswith('Heading 1'):
            if current_part_name != "":
                # saving part
                i += 1
                current_document.save(f"{output_dir}/{basename}-part-{i}-{current_part_name}.docx")

            current_document = Document()
            current_part_name = p.text
            get_para_data(current_document, p)
        else:
            get_para_data(current_document, p)

if __name__ == "__main__":
    split_word("test/report.docx")
